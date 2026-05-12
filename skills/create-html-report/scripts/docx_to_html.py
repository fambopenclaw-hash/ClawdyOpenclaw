#!/usr/bin/env python3
"""
docx_to_html.py — Read a .docx file and output document content as structured JSON.
Usage: python3 docx_to_html.py <file.docx>
"""
import json, sys

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: uv pip install python-docx --python /tmp/reportenv/bin/python3")
    sys.exit(1)

doc = Document(sys.argv[1])
paragraphs_data = []

for para in doc.paragraphs:
    if para.text.strip():
        level = para.style.name if para.style else "Normal"
        paragraphs_data.append({"style": level, "text": para.text.strip()})

tables_data = []
for table in doc.tables:
    rows_data = [[cell.text for cell in row.cells] for row in table.rows]
    tables_data.append(rows_data)

result = {
    "num_paragraphs": len(paragraphs_data),
    "num_tables": len(tables_data),
    "paragraphs": paragraphs_data,
    "tables": tables_data
}
print(json.dumps(result, indent=2, default=str))